import pytest
import io
from unittest.mock import patch, MagicMock
from fastapi.testclient import TestClient
from app import app, extract_text_from_url, extract_text_from_pdf, extract_text_from_docx, extract_text_from_txt, summarize_text
import docx
import PyPDF2

client = TestClient(app)

def test_extract_text_from_txt():
    content = b"Hello world! This is a test."
    text = extract_text_from_txt(content)
    assert text == "Hello world! This is a test."

def test_summarize_text():
    # A short text shouldn't be summarized, or handles it well
    short_text = "This is short."
    assert summarize_text(short_text) == short_text

    # A longer text
    long_text = "This is the first sentence. This is the second sentence. The third sentence is also here. The fourth sentence adds length. The fifth sentence completes it. This paragraph provides some context. Here is another sentence." * 3
    summary = summarize_text(long_text, per=0.3)
    assert len(summary) > 0

@patch('app.requests.get')
def test_extract_text_from_url(mock_get):
    # Mock response
    mock_response = MagicMock()
    mock_response.content = b"<html><body><p>Test paragraph 1</p><p>Test paragraph 2</p></body></html>"
    mock_response.raise_for_status.return_value = None
    mock_get.return_value = mock_response

    text = extract_text_from_url("http://example.com")
    assert "Test paragraph 1" in text
    assert "Test paragraph 2" in text

def test_extract_text_from_docx():
    # Create a real docx in memory
    doc = docx.Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    content = file_stream.read()

    text = extract_text_from_docx(content)
    assert "First paragraph." in text
    assert "Second paragraph." in text

def test_extract_text_from_pdf():
    # It's hard to dynamically create a PDF, so we mock PyPDF2.PdfReader
    with patch('app.PyPDF2.PdfReader') as MockPdfReader:
        mock_reader_instance = MagicMock()
        mock_page = MagicMock()
        mock_page.extract_text.return_value = "This is a PDF test."
        mock_reader_instance.pages = [mock_page]
        MockPdfReader.return_value = mock_reader_instance

        # We can pass any dummy content since it's mocked
        text = extract_text_from_pdf(b"dummy pdf content")
        assert "This is a PDF test." in text

@patch('app.requests.get')
def test_api_summarize_url(mock_get):
    mock_response = MagicMock()
    mock_response.content = b"<html><body><p>The quick brown fox jumps over the lazy dog. The quick brown fox jumps over the lazy dog. The quick brown fox jumps over the lazy dog. The quick brown fox jumps over the lazy dog.</p></body></html>"
    mock_response.raise_for_status.return_value = None
    mock_get.return_value = mock_response

    response = client.post("/summarize-url", data={"url": "http://example.com", "summary_ratio": 0.5})
    assert response.status_code == 200
    data = response.json()
    assert "extracted_text" in data
    assert "summary" in data
    assert "The quick brown fox" in data["extracted_text"]

def test_api_summarize_file_txt():
    file_content = b"The quick brown fox jumps over the lazy dog. The quick brown fox jumps over the lazy dog. The quick brown fox jumps over the lazy dog. The quick brown fox jumps over the lazy dog."

    response = client.post(
        "/summarize-file",
        data={"summary_ratio": 0.5},
        files={"file": ("test.txt", file_content, "text/plain")}
    )
    assert response.status_code == 200
    data = response.json()
    assert data["filename"] == "test.txt"
    assert "extracted_text" in data
    assert "summary" in data
    assert "The quick brown fox" in data["extracted_text"]
